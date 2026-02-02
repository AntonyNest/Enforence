"""
Експорт ТЗ у формат DOCX.

Генерація Word документу згідно структури КМУ №205.
"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from src.utils.logger import get_logger

logger = get_logger(__name__)


def create_tz_document(sections: list[dict[str, Any]], metadata: dict[str, Any]) -> Document:
    """
    Створення DOCX документу ТЗ.

    Args:
        sections: Список секцій ТЗ з контентом.
        metadata: Метадані проєкту (назва, замовник, тощо).

    Returns:
        python-docx Document об'єкт.
    """
    doc = Document()

    # Налаштування стилів документу
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # Налаштування полів сторінки
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Титульна сторінка
    _add_title_page(doc, metadata)

    # Секції ТЗ
    for section_data in sections:
        _add_section(doc, section_data)

    logger.info(
        "docx_document_created",
        sections_count=len(sections),
        project_name=metadata.get("name", ""),
    )

    return doc


def _add_title_page(doc: Document, metadata: dict[str, Any]) -> None:
    """Додавання титульної сторінки."""
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ЗАТВЕРДЖУЮ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")
    doc.add_paragraph("")

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("ТЕХНІЧНЕ ЗАВДАННЯ")
    run.bold = True
    run.font.size = Pt(16)

    if name := metadata.get("name"):
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"на створення {name}")
        run.font.size = Pt(14)

    doc.add_page_break()


def _add_section(doc: Document, section_data: dict[str, Any]) -> None:
    """Додавання секції ТЗ до документу."""
    section_id = section_data.get("id", "")
    title = section_data.get("title", "")

    # Заголовок секції
    heading = doc.add_heading(f"{section_id}. {title}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Контент секції
    if content := section_data.get("content"):
        paragraph = doc.add_paragraph(content)
        paragraph.style = doc.styles["Normal"]

    # Підсекції
    for subsection in section_data.get("subsections", []):
        sub_id = subsection.get("id", "")
        sub_title = subsection.get("title", "")

        sub_heading = doc.add_heading(f"{sub_id}. {sub_title}", level=2)
        sub_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if sub_content := subsection.get("content"):
            paragraph = doc.add_paragraph(sub_content)
            paragraph.style = doc.styles["Normal"]


def save_document(doc: Document, output_path: Path) -> Path:
    """
    Збереження DOCX документу на диск.

    Args:
        doc: python-docx Document об'єкт.
        output_path: Шлях для збереження файлу.

    Returns:
        Шлях до збереженого файлу.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("docx_saved", path=str(output_path))
    return output_path
